# docx_parser.py
"""Wordファイル(.docx/.doc)から就業規則を抽出するモジュール"""

import os
import re
import subprocess
import tempfile
from typing import Dict, List
from docx import Document
from structure_analyzer import StructureAnalyzer


class DocxParser:
    """Wordファイルから就業規則を抽出するクラス"""

    def __init__(self):
        self.structure_analyzer = StructureAnalyzer()

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """.docファイルをLibreOfficeで.docxに変換"""
        print(f".doc形式を検出、.docxに変換中...")
        output_dir = tempfile.mkdtemp()
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'docx',
            '--outdir', output_dir, doc_path
        ], check=True, timeout=60)

        # 変換後のファイルパスを取得
        basename = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(output_dir, basename + '.docx')

        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"変換後のファイルが見つかりません: {docx_path}")

        print(f".docx変換完了")
        return docx_path

    def extract_from_docx(self, file_path: str) -> Dict:
        """Wordファイル(.docx/.doc)から就業規則を抽出"""
        converted_path = None
        try:
            # .docの場合は.docxに変換
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.doc':
                converted_path = self._convert_doc_to_docx(file_path)
                target_path = converted_path
            else:
                target_path = file_path

            doc = Document(target_path)
            print(f"Wordファイルを読み込み中: {target_path}")

            # 段落からテキスト抽出
            paragraphs_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs_text.append(text)

            # 表からテキスト抽出
            detected_tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    detected_tables.append({
                        'page': i + 1,
                        'data': table_data
                    })
                    paragraphs_text.append("\n[表データ検出]")
                    paragraphs_text.append("-" * 50)
                    for row_data in table_data:
                        paragraphs_text.append(" | ".join(row_data))
                    paragraphs_text.append("-" * 50)

            full_text = "\n".join(paragraphs_text)
            print(f"Wordから{len(full_text)}文字抽出しました")

            # 会社情報を抽出
            company_info = self._extract_company_info(full_text)

            # 構造解析
            print("構造解析を実行中...")
            structure_result = self.structure_analyzer.analyze(full_text)
            print(f"  章: {structure_result['metadata']['total_chapters']}個")
            print(f"  条: {structure_result['metadata']['total_articles']}個")
            print(f"  項: {structure_result['metadata']['total_items']}個")

            return {
                "success": True,
                "company_info": company_info,
                "raw_text": full_text,
                "structure": structure_result["structure"],
                "tables": detected_tables
            }

        except Exception as e:
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e)
            }
        finally:
            # 変換した一時ファイルを削除
            if converted_path and os.path.exists(converted_path):
                try:
                    os.remove(converted_path)
                    os.rmdir(os.path.dirname(converted_path))
                except:
                    pass

    def _extract_company_info(self, text: str) -> Dict:
        """会社情報を抽出"""
        info = {
            "company_name": None,
            "address": None,
            "effective_date": None
        }

        company_patterns = [
            r'([^\s]+(?:株式会社|有限会社|合同会社|一般社団法人|NPO法人))',
            r'(株式会社[^\s]+|有限会社[^\s]+|合同会社[^\s]+)',
        ]
        for pattern in company_patterns:
            match = re.search(pattern, text[:500])
            if match:
                info["company_name"] = match.group(1)
                break

        date_pattern = r'(令和|平成)(\d+)年(\d+)月(\d+)日'
        match = re.search(date_pattern, text[:500])
        if match:
            info["effective_date"] = match.group(0)

        return info
